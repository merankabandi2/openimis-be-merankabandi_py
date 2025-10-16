import os
import json
from datetime import datetime, date
from decimal import Decimal
from django.db.models import Count, Sum, Q, F
from django.utils import timezone
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from location.models import Location
from social_protection.models import GroupBeneficiary, BenefitPlan
from individual.models import Individual, GroupIndividual, Group
from payroll.models import BenefitConsumption
from .models import (
    Section, Indicator, IndicatorAchievement, 
    ResultFrameworkSnapshot, IndicatorCalculationRule,
    MonetaryTransfer, SensitizationTraining, 
    BehaviorChangePromotion, MicroProject
)

# Host communes for refugee/host community separation
HOST_COMMUNES = ['Butezi', 'Ruyigi', 'Kiremba', 'Gasorwe', 'Gashoho', 'Muyinga', 'Cankuzo']

# Refugee collines/camps for refugee/host community separation
REFUGEE_COLLINES = []


class ResultFrameworkService:
    """Service for result framework calculations and document generation"""
    
    def __init__(self):
        self.calculation_methods = {
            # Development indicators (sections 1-3)
            'count_households_registered': self._count_households_registered,
            'count_households_refugees': self._count_households_refugees,
            'count_households_host': self._count_households_host,
            'count_beneficiaries_social_protection': self._count_beneficiaries_social_protection,
            'count_beneficiaries_women': self._count_beneficiaries_women,
            'count_beneficiaries_unconditional_transfers': self._count_beneficiaries_unconditional_transfers,
            'count_beneficiaries_emergency_transfers': self._count_beneficiaries_emergency_transfers,
            'count_beneficiaries_refugees': self._count_beneficiaries_refugees,
            'count_beneficiaries_host_communities': self._count_beneficiaries_host_communities,
            'count_beneficiaries_employment': self._count_beneficiaries_employment,
            'count_beneficiaries_employment_women': self._count_beneficiaries_employment_women,
            'count_beneficiaries_employment_refugees': self._count_beneficiaries_employment_refugees,
            'count_beneficiaries_employment_host': self._count_beneficiaries_employment_host,
            'count_farmers_received_services': self._count_farmers_received_services,
            
            # Intermediate indicators (sections 4-8)
            'count_provinces_with_transfers': self._count_provinces_with_transfers,
            'calculate_payment_timeliness': self._calculate_payment_timeliness,
            'calculate_behavior_change_participation': self._calculate_behavior_change_participation,
            'count_approved_business_plans': self._count_approved_business_plans,
            'count_approved_business_plans_women': self._count_approved_business_plans_women,
            'count_approved_business_plans_batwa': self._count_approved_business_plans_batwa,
            'count_climate_resilient_activities': self._count_climate_resilient_activities,
            'calculate_digital_payment_percentage': self._calculate_digital_payment_percentage,
        }
    
    def calculate_indicator_value(self, indicator_id, date_from=None, date_to=None, location=None):
        """Calculate indicator value based on its configuration"""
        try:
            indicator = Indicator.objects.get(id=indicator_id)
            rule = IndicatorCalculationRule.objects.filter(indicator=indicator, is_active=True).first()
            
            if not rule:
                # Default to manual if no rule exists
                return self._get_latest_achievement(indicator, date_from, date_to)
            
            if rule.calculation_type == 'MANUAL':
                return self._get_latest_achievement(indicator, date_from, date_to)
            
            elif rule.calculation_type == 'SYSTEM':
                method_name = rule.calculation_method
                if method_name in self.calculation_methods:
                    return self.calculation_methods[method_name](
                        indicator, date_from, date_to, location, rule.calculation_config
                    )
                else:
                    return {'value': 0, 'error': f'Unknown calculation method: {method_name}'}
            
            elif rule.calculation_type == 'MIXED':
                # Get system calculated value
                system_value = 0
                if rule.calculation_method in self.calculation_methods:
                    system_result = self.calculation_methods[rule.calculation_method](
                        indicator, date_from, date_to, location, rule.calculation_config
                    )
                    system_value = system_result.get('value', 0)
                
                # Get manual adjustment
                manual_result = self._get_latest_achievement(indicator, date_from, date_to)
                manual_value = manual_result.get('value', 0)
                
                # Combine based on config
                combine_method = rule.calculation_config.get('combine_method', 'add')
                if combine_method == 'add':
                    final_value = system_value + manual_value
                elif combine_method == 'max':
                    final_value = max(system_value, manual_value)
                elif combine_method == 'replace':
                    final_value = manual_value if manual_value > 0 else system_value
                else:
                    final_value = system_value
                
                return {
                    'value': final_value,
                    'system_value': system_value,
                    'manual_value': manual_value,
                    'calculation_type': 'MIXED'
                }
                
        except Exception as e:
            return {'value': 0, 'error': str(e)}
    
    def _get_latest_achievement(self, indicator, date_from=None, date_to=None):
        """Get the latest manual achievement entry"""
        query = IndicatorAchievement.objects.filter(indicator=indicator)
        
        if date_from:
            query = query.filter(date__gte=date_from)
        if date_to:
            query = query.filter(date__lte=date_to)
        
        latest = query.order_by('-date', '-timestamp').first()
        
        if latest:
            return {
                'value': float(latest.achieved),
                'date': latest.date,
                'comment': latest.comment,
                'calculation_type': 'MANUAL'
            }
        return {'value': 0, 'calculation_type': 'MANUAL'}
    
    # Calculation methods for each indicator type
    def _count_households_registered(self, indicator, date_from, date_to, location, config):
        """Count total households registered (Indicator 1)"""
        query = Group.objects.filter(
            is_deleted=False
        )
        
        if date_from:
            query = query.filter(date_created__gte=date_from)
        if date_to:
            query = query.filter(date_created__lte=date_to)
        if location:
            query = query.filter(location__parent__parent=location)
        
        count = query.count()
        return {'value': count, 'calculation_type': 'SYSTEM'}
    
    def _count_households_refugees(self, indicator, date_from, date_to, location, config):
        """Count refugee households (Indicator 2)"""
        query = Group.objects.filter(
            is_deleted=False
        ).exclude(
            location__parent__name__in=HOST_COMMUNES
        )
        
        if date_from:
            query = query.filter(date_created__gte=date_from)
        if date_to:
            query = query.filter(date_created__lte=date_to)
        if location:
            query = query.filter(location__parent__parent=location)
        
        count = query.count()
        
        # Get gender breakdown
        gender_data = {}
        for gb in query.select_related('head'):
            if gb.group.head and gb.group.head.json_ext:
                gender = gb.group.head.json_ext.get('sexe', 'unknown')
                gender_data[gender] = gender_data.get(gender, 0) + 1
        
        return {
            'value': count,
            'gender_breakdown': gender_data,
            'calculation_type': 'SYSTEM'
        }
    
    def _count_households_host(self, indicator, date_from, date_to, location, config):
        """Count host community households (Indicator 3)"""
        query = Group.objects.filter(
            is_deleted=False,
            location__parent__name__in=HOST_COMMUNES
        )
        
        if date_from:
            query = query.filter(date_created__gte=date_from)
        if date_to:
            query = query.filter(date_created__lte=date_to)
        if location:
            query = query.filter(location__parent__parent=location)

        count = query.count()
        return {'value': count, 'calculation_type': 'SYSTEM'}
    
    def _count_beneficiaries_social_protection(self, indicator, date_from, date_to, location, config):
        """Count total beneficiaries (Indicator 5)"""
        query = GroupBeneficiary.objects.filter(
            is_deleted=False,
            status__in=['ACTIVE', 'VALIDATED', 'POTENTIAL']
        )
        
        if date_from:
            query = query.filter(date_created__gte=date_from)
        if date_to:
            query = query.filter(date_created__gte=date_to)
        if location:
            query = query.filter(group__location__parent__parent=location)
        
        count = query.distinct().count()
        return {'value': count, 'calculation_type': 'SYSTEM'}
    
    def _count_beneficiaries_women(self, indicator, date_from, date_to, location, config):
        """Count female beneficiaries (Indicator 6)"""
        query = GroupBeneficiary.objects.filter(
            is_deleted=False,
            status__in=['ACTIVE', 'VALIDATED', 'POTENTIAL'],
            group__groupindividuals__individual__json_ext__sexe='F',
            group__groupindividuals__recipient_type='PRIMARY'
        )
        
        if date_from:
            query = query.filter(date_due__gte=date_from)
        if date_to:
            query = query.filter(date_due__lte=date_to)
        if location:
            query = query.filter(group__location__parent__parent=location)

        count = query.distinct().count()
        return {'value': count, 'calculation_type': 'SYSTEM'}
    
    def _count_beneficiaries_unconditional_transfers(self, indicator, date_from, date_to, location, config):
        """Count beneficiaries of unconditional transfers (Indicator 7)"""
        # Filter by benefit plan code for unconditional transfers
        query = GroupBeneficiary.objects.filter(
            is_deleted=False,
            status__in=['ACTIVE', 'VALIDATED', 'POTENTIAL'],
            group__groupindividuals__individual__json_ext__sexe='F',
            group__groupindividuals__recipient_type='PRIMARY',
            benefit_plan__code__in=['1.2']  # Adjust based on actual codes
        )
        
        if date_from:
            query = query.filter(date_due__gte=date_from)
        if date_to:
            query = query.filter(date_due__lte=date_to)
        if location:
            query = query.filter(group__location__parent__parent=location)
        
        count = query.distinct().count()
        return {'value': count, 'calculation_type': 'SYSTEM'}
    
    def _count_beneficiaries_employment(self, indicator, date_from, date_to, location, config):
        """Count beneficiaries of employment interventions (Indicator 11)"""
        # Count from training and microproject participants
        training_query = SensitizationTraining.objects.filter(validation_status='VALIDATED')
        microproject_query = MicroProject.objects.filter(validation_status='VALIDATED')
        
        if date_from:
            # training_query = training_query.filter(sensitization_date__gte=date_from)
            microproject_query = microproject_query.filter(report_date__gte=date_from)
        if date_to:
            # training_query = training_query.filter(sensitization_date__lte=date_to)
            microproject_query = microproject_query.filter(report_date__lte=date_to)
        if location:
            # training_query = training_query.filter(location__parent__parent=location)
            microproject_query = microproject_query.filter(location__parent__parent=location)
        
        # Sum participants
        # training_total = training_query.aggregate(
        #     total=Sum('male_participants') + Sum('female_participants')
        # )['total'] or 0
        
        microproject_total = microproject_query.aggregate(
            total=Sum('male_participants') + Sum('female_participants')
        )['total'] or 0

        return {'value': microproject_total, 'calculation_type': 'SYSTEM'}
    
    def _count_provinces_with_transfers(self, indicator, date_from, date_to, location, config):
        """Count provinces implementing transfers (Indicator 16)"""
        query = MonetaryTransfer.objects.all()
        
        if date_from:
            query = query.filter(transfer_date__gte=date_from)
        if date_to:
            query = query.filter(transfer_date__lte=date_to)
        
        # Get unique provinces
        provinces = query.values('location__parent__parent').distinct().count()
        return {'value': provinces, 'calculation_type': 'SYSTEM'}
    
    def _calculate_payment_timeliness(self, indicator, date_from, date_to, location, config):
        """Calculate percentage of beneficiaries paid on time (Indicator 17)"""
        # This would need payment schedule data to calculate properly
        # For now, return manual value
        return self._get_latest_achievement(indicator, date_from, date_to)
    
    def _count_beneficiaries_emergency_transfers(self, indicator, date_from, date_to, location, config):
        query = GroupBeneficiary.objects.filter(
            is_deleted=False,
            status__in=['ACTIVE', 'VALIDATED', 'POTENTIAL'],
            group__groupindividuals__individual__json_ext__sexe='F',
            group__groupindividuals__recipient_type='PRIMARY',
            benefit_plan__code__in=['1.1']  # Adjust based on actual codes
        )
        
        if date_from:
            query = query.filter(date_due__gte=date_from)
        if date_to:
            query = query.filter(date_due__lte=date_to)
        if location:
            query = query.filter(group__location__parent__parent=location)
        
        count = query.distinct().count()
        return {'value': count, 'calculation_type': 'SYSTEM'}
    
    def _count_beneficiaries_refugees(self, indicator, date_from, date_to, location, config):
        query = GroupBeneficiary.objects.filter(
            is_deleted=False,
            status__in=['ACTIVE', 'VALIDATED', 'POTENTIAL'],
            group__groupindividuals__individual__json_ext__sexe='F',
            group__groupindividuals__recipient_type='PRIMARY',
            benefit_plan__code__in=['1.4']
        )
        
        if date_from:
            query = query.filter(date_due__gte=date_from)
        if date_to:
            query = query.filter(date_due__lte=date_to)
        if location:
            query = query.filter(group__location__parent__parent=location)

        count = query.distinct().count()
        return {'value': count, 'calculation_type': 'SYSTEM'}
        return self._get_latest_achievement(indicator, date_from, date_to)
    
    def _count_beneficiaries_host_communities(self, indicator, date_from, date_to, location, config):
        query = GroupBeneficiary.objects.filter(
            is_deleted=False,
            status__in=['ACTIVE', 'VALIDATED', 'POTENTIAL'],
            group__groupindividuals__individual__json_ext__sexe='F',
            group__groupindividuals__recipient_type='PRIMARY',
            group__location__parent__name__in=HOST_COMMUNES
        )
        
        if date_from:
            query = query.filter(date_created__gte=date_from)
        if date_to:
            query = query.filter(date_created__lte=date_to)
        if location:
            query = query.filter(group__location__parent__parent=location)

        count = query.count()
        return {'value': count, 'calculation_type': 'SYSTEM'}
    
    def _count_beneficiaries_employment_women(self, indicator, date_from, date_to, location, config):
        # Count from microproject participants
        microproject_query = MicroProject.objects.filter(validation_status='VALIDATED')
        
        if date_from:
            microproject_query = microproject_query.filter(report_date__gte=date_from)
        if date_to:
            microproject_query = microproject_query.filter(report_date__lte=date_to)
        if location:
            microproject_query = microproject_query.filter(location__parent__parent=location)
        
        # Sum participants
        microproject_total = microproject_query.aggregate(
            total=Sum('female_participants')
        )['total'] or 0

        return {'value': microproject_total, 'calculation_type': 'SYSTEM'}
    
    def _count_beneficiaries_employment_refugees(self, indicator, date_from, date_to, location, config):
        # Count from microproject participants
        microproject_query = MicroProject.objects.filter(validation_status='VALIDATED')
        
        if date_from:
            microproject_query = microproject_query.filter(report_date__gte=date_from)
        if date_to:
            microproject_query = microproject_query.filter(report_date__lte=date_to)
        if location:
            microproject_query = microproject_query.filter(
            location__name__in=REFUGEE_COLLINES)

        # Sum participants
        microproject_total = microproject_query.aggregate(
            total=Sum('male_participants') + Sum('female_participants')
        )['total'] or 0

        return {'value': microproject_total, 'calculation_type': 'SYSTEM'}
    
    def _count_beneficiaries_employment_host(self, indicator, date_from, date_to, location, config):
        # Count from microproject participants
        microproject_query = MicroProject.objects.filter(validation_status='VALIDATED')
        
        if date_from:
            microproject_query = microproject_query.filter(report_date__gte=date_from)
        if date_to:
            microproject_query = microproject_query.filter(report_date__lte=date_to)
        if location:
            microproject_query = microproject_query.filter(
            location__parent__name__in=HOST_COMMUNES)
        
        # Sum participants
        microproject_total = microproject_query.aggregate(
            total=Sum('male_participants') + Sum('female_participants')
        )['total'] or 0

        return {'value': microproject_total, 'calculation_type': 'SYSTEM'}

    def _count_farmers_received_services(self, indicator, date_from, date_to, location, config):
        # Count from microproject participants
        microproject_query = MicroProject.objects.filter(validation_status='VALIDATED')
        
        if date_from:
            microproject_query = microproject_query.filter(report_date__gte=date_from)
        if date_to:
            microproject_query = microproject_query.filter(report_date__lte=date_to)
        if location:
            microproject_query = microproject_query.filter(location__parent__parent=location)
        
        # Sum participants
        microproject_total = microproject_query.aggregate(
            total=Sum('agriculture_beneficiaries')
        )['total'] or 0

        return {'value': microproject_total, 'calculation_type': 'SYSTEM'}
    
    def _calculate_behavior_change_participation(self, indicator, date_from, date_to, location, config):
        """Calculate percentage of beneficiaries participating in behavior change activities (Indicator 18)"""
        # Count beneficiaries participating in sensitization trainings
        training_query = SensitizationTraining.objects.filter(validation_status='VALIDATED')
        
        if date_from:
            training_query = training_query.filter(report_date__gte=date_from)
        if date_to:
            training_query = training_query.filter(report_date__lte=date_to)
        if location:
            training_query = training_query.filter(location__parent__parent=location)
        
        # Get total participants
        total_participants = training_query.aggregate(
            total=Sum('male_participants') + Sum('female_participants')
        )['total'] or 0
        
        # Get total beneficiaries in the area
        beneficiary_query = GroupBeneficiary.objects.filter(
            is_deleted=False,
            status__in=['ACTIVE', 'VALIDATED']
        )
        if location:
            beneficiary_query = beneficiary_query.filter(group__location__parent__parent=location)
        
        total_beneficiaries = beneficiary_query.count()
        
        if total_beneficiaries > 0:
            percentage = (total_participants / total_beneficiaries) * 100
            return {'value': min(percentage, 100), 'calculation_type': 'SYSTEM'}
        
        return {'value': 0, 'calculation_type': 'SYSTEM'}
    
    def _count_approved_business_plans(self, indicator, date_from, date_to, location, config):
        """Count beneficiaries with approved business plans (Indicator 20)"""
        # Count from microprojects with approved status
        query = MicroProject.objects.filter(
            validation_status='VALIDATED'
        )
        
        if date_from:
            query = query.filter(report_date__gte=date_from)
        if date_to:
            query = query.filter(report_date__lte=date_to)
        if location:
            query = query.filter(location__parent__parent=location)
        
        # Sum all participants
        total = query.aggregate(
            total=Sum('male_participants') + Sum('female_participants')
        )['total'] or 0
        
        return {'value': total, 'calculation_type': 'MIXED'}
    
    def _count_approved_business_plans_women(self, indicator, date_from, date_to, location, config):
        """Count female beneficiaries with approved business plans (Indicator 21)"""
        query = MicroProject.objects.filter(
            validation_status='VALIDATED'
        )
        
        if date_from:
            query = query.filter(report_date__gte=date_from)
        if date_to:
            query = query.filter(report_date__lte=date_to)
        if location:
            query = query.filter(location__parent__parent=location)
        
        total = query.aggregate(total=Sum('female_participants'))['total'] or 0
        
        return {'value': total, 'calculation_type': 'MIXED'}
    
    def _count_approved_business_plans_batwa(self, indicator, date_from, date_to, location, config):
        """Count Batwa beneficiaries with approved business plans (Indicator 22)"""
        # This would need specific tracking of Batwa beneficiaries
        # For now, use manual entry
        return self._get_latest_achievement(indicator, date_from, date_to)
    
    def _count_climate_resilient_activities(self, indicator, date_from, date_to, location, config):
        """Count climate-resilient productive activities (Indicator 23)"""
        query = MicroProject.objects.filter(
            validation_status='VALIDATED'
        )
        
        if date_from:
            query = query.filter(report_date__gte=date_from)
        if date_to:
            query = query.filter(report_date__lte=date_to)
        if location:
            query = query.filter(location__parent__parent=location)
        
        count = query.count()
        return {'value': count, 'calculation_type': 'MIXED'}
    
    def _calculate_digital_payment_percentage(self, indicator, date_from, date_to, location, config):
        """Calculate percentage of beneficiaries receiving digital payments (Indicator 28)"""
        # Count beneficiaries with digital payment method
        query = BenefitConsumption.objects.filter(
            individual__is_deleted=False,
            json_ext__payment_method='DIGITAL'
        )
        
        if date_from:
            query = query.filter(date_created__gte=date_from)
        if date_to:
            query = query.filter(date_created__lte=date_to)
        if location:
            query = query.filter(individual__group__location__parent__parent=location)
        
        digital_count = query.values('individual').distinct().count()
        
        # Get total beneficiaries who received payments
        total_query = BenefitConsumption.objects.filter(
            individual__is_deleted=False
        )
        
        if date_from:
            total_query = total_query.filter(date_created__gte=date_from)
        if date_to:
            total_query = total_query.filter(date_created__lte=date_to)
        if location:
            total_query = total_query.filter(individual__group__location__parent__parent=location)
        
        total_count = total_query.values('individual').distinct().count()
        
        if total_count > 0:
            percentage = (digital_count / total_count) * 100
            return {'value': percentage, 'calculation_type': 'SYSTEM'}
        
        return {'value': 0, 'calculation_type': 'SYSTEM'}
    
    def create_snapshot(self, name, description, user, date_from=None, date_to=None):
        """Create a complete snapshot of the result framework"""
        snapshot_data = {
            'sections': [],
            'metadata': {
                'created_date': timezone.now().isoformat(),
                'created_by': user.username if user else 'System',
                'date_from': date_from.isoformat() if date_from else None,
                'date_to': date_to.isoformat() if date_to else None,
            }
        }
        
        for section in Section.objects.all().prefetch_related('indicators'):
            section_data = {
                'id': section.id,
                'name': section.name,
                'indicators': []
            }
            
            for indicator in section.indicators.all():
                # Calculate current value
                result = self.calculate_indicator_value(
                    indicator.id, 
                    date_from=date_from, 
                    date_to=date_to
                )
                print([indicator.name, result])
                achieved_value = result.get('value', 0)
                target_value = float(indicator.target) if indicator.target else 0
                
                indicator_data = {
                    'id': indicator.id,
                    'name': indicator.name,
                    'pbc': indicator.pbc or '',
                    'baseline': float(indicator.baseline) if indicator.baseline else 0,
                    'target': target_value,
                    'achieved': achieved_value,
                    'percentage': (achieved_value / target_value * 100) if target_value > 0 else 0,
                    'calculation_type': result.get('calculation_type', 'MANUAL'),
                    'observation': indicator.observation or ''
                }

                # Save IndicatorAchievement record if value was calculated (not manual)
                if result.get('calculation_type') in ['SYSTEM', 'MIXED'] and achieved_value > 0:
                    achievement_date = date_to if date_to else timezone.now().date()
 
                    # Create or update achievement for this date
                    achievement = IndicatorAchievement.objects.create(
                        indicator=indicator,
                        date=achievement_date,
                        achieved=Decimal(str(achieved_value)),
                        comment=f'Auto-generated from snapshot: {name} (Calculation: {result.get("calculation_type")})'
                    )

                # Add any additional data from calculation
                if 'gender_breakdown' in result:
                    indicator_data['gender_breakdown'] = result['gender_breakdown']
                if 'error' in result:
                    indicator_data['error'] = result['error']
                
                section_data['indicators'].append(indicator_data)
            
            snapshot_data['sections'].append(section_data)
        
        # Create snapshot record
        snapshot = ResultFrameworkSnapshot.objects.create(
            name=name,
            description=description,
            created_by=user,
            data=snapshot_data,
            status='DRAFT'
        )
        
        return snapshot
    
    def generate_document(self, snapshot_id=None, format='docx'):
        """Generate result framework document"""
        if snapshot_id:
            snapshot = ResultFrameworkSnapshot.objects.get(id=snapshot_id)
            data = snapshot.data
        else:
            # Generate current data
            temp_snapshot = self.create_snapshot("Temporary", "", None)
            data = temp_snapshot.data
            temp_snapshot.delete()  # Clean up
        
        if format == 'docx':
            return self._generate_docx(data)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _generate_docx(self, data):
        """Generate DOCX document from snapshot data"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading('CADRE DE RÉSULTATS - MERANKABANDI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Généré le: {data['metadata']['created_date'][:10]}")
        
        if data['metadata'].get('date_from') or data['metadata'].get('date_to'):
            period_para = doc.add_paragraph()
            period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            period_text = "Période: "
            if data['metadata'].get('date_from'):
                period_text += f"Du {data['metadata']['date_from'][:10]} "
            if data['metadata'].get('date_to'):
                period_text += f"Au {data['metadata']['date_to'][:10]}"
            period_para.add_run(period_text)
        
        doc.add_paragraph()  # Empty line
        
        # Add sections and indicators
        for section_data in data['sections']:
            # Section header
            doc.add_heading(section_data['name'], level=1)
            
            # Create table for indicators
            if section_data['indicators']:
                # Table with 7 columns
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header row
                headers = ['Indicateur', 'PBC', 'Baseline', 'Cible', 'Réalisé', '% Achevé', 'Observation']
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    # Bold headers
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add indicator rows
                for indicator in section_data['indicators']:
                    row = table.add_row()
                    cells = row.cells
                    
                    cells[0].text = indicator['name']
                    cells[1].text = indicator.get('pbc', '')
                    cells[2].text = f"{indicator['baseline']:,.0f}"
                    cells[3].text = f"{indicator['target']:,.0f}"
                    cells[4].text = f"{indicator['achieved']:,.0f}"
                    cells[5].text = f"{indicator['percentage']:.1f}%"
                    cells[6].text = indicator.get('observation', '')
                    
                    # Color code based on percentage
                    percentage = indicator['percentage']
                    if percentage >= 80:
                        color = RGBColor(0, 128, 0)  # Green
                    elif percentage >= 50:
                        color = RGBColor(255, 165, 0)  # Orange
                    else:
                        color = RGBColor(255, 0, 0)  # Red
                    
                    # Apply color to percentage cell
                    for paragraph in cells[5].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = color
                
                # Adjust column widths
                for i, width in enumerate([3.5, 0.8, 0.8, 0.8, 0.8, 0.8, 1.5]):
                    for cell in table.columns[i].cells:
                        cell.width = Inches(width)
            
            doc.add_paragraph()  # Space between sections
        
        # Summary statistics
        doc.add_page_break()
        doc.add_heading('RÉSUMÉ', level=1)
        
        total_indicators = sum(len(s['indicators']) for s in data['sections'])
        achieved_80_plus = sum(1 for s in data['sections'] for i in s['indicators'] if i['percentage'] >= 80)
        achieved_50_plus = sum(1 for s in data['sections'] for i in s['indicators'] if 50 <= i['percentage'] < 80)
        achieved_below_50 = sum(1 for s in data['sections'] for i in s['indicators'] if i['percentage'] < 50)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Nombre total d'indicateurs: {total_indicators}\n")
        summary_para.add_run(f"Indicateurs avec ≥80% d'achèvement: {achieved_80_plus}\n")
        summary_para.add_run(f"Indicateurs avec 50-79% d'achèvement: {achieved_50_plus}\n")
        summary_para.add_run(f"Indicateurs avec <50% d'achèvement: {achieved_below_50}\n")
        
        return doc